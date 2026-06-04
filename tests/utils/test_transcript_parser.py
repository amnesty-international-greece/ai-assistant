"""Tests for the transcript parser utility."""

from pathlib import Path

import pytest

from src.utils.transcript_parser import parse_transcript


def test_parse_vtt(tmp_path):
    """Parses WebVTT with timestamps, speaker labels, and collapse."""
    vtt = tmp_path / "meeting.vtt"
    vtt.write_text(
        "WEBVTT\n\n"
        "1\n"
        "00:00:01.000 --> 00:00:05.000\n"
        "Speaker A: Hello everyone.\n\n"
        "2\n"
        "00:00:05.500 --> 00:00:10.000\n"
        "Speaker A: Let's begin.\n\n"
        "3\n"
        "00:00:10.500 --> 00:00:15.000\n"
        "Speaker B: I have a question.\n",
        encoding="utf-8",
    )
    result = parse_transcript(vtt)
    # Speaker A's consecutive lines should be collapsed
    assert "Speaker A: Hello everyone. Let's begin." in result
    assert "Speaker B: I have a question." in result
    # No timestamps or sequence numbers
    assert "-->" not in result
    assert "WEBVTT" not in result


def test_parse_txt(tmp_path):
    """Returns plain text as-is."""
    txt = tmp_path / "notes.txt"
    txt.write_text("Σημειώσεις Γ.Γ.\nΘέμα 1: Εγκριση", encoding="utf-8")
    result = parse_transcript(txt)
    assert "Σημειώσεις Γ.Γ." in result
    assert "Θέμα 1: Εγκριση" in result


def test_parse_docx(tmp_path):
    """Parses a .docx file (requires python-docx)."""
    pytest.importorskip("docx")  # skip if python-docx not installed
    from docx import Document

    docx_path = tmp_path / "transcript.docx"
    doc = Document()
    doc.add_paragraph("Παρόντες: Α, Β, Γ")
    doc.add_paragraph("Θέμα 1: Συζήτηση")
    doc.save(str(docx_path))

    result = parse_transcript(docx_path)
    assert "Παρόντες" in result
    assert "Θέμα 1" in result


def test_unsupported_extension(tmp_path):
    """Raises ValueError for unsupported file types."""
    bad = tmp_path / "meeting.mp4"
    bad.write_bytes(b"fake")
    with pytest.raises(ValueError, match="Unsupported"):
        parse_transcript(bad)


def test_file_not_found():
    """Raises FileNotFoundError for missing files."""
    with pytest.raises(FileNotFoundError):
        parse_transcript("/nonexistent/path/transcript.vtt")


def test_parse_vtt_empty(tmp_path):
    """Handles empty VTT gracefully."""
    vtt = tmp_path / "empty.vtt"
    vtt.write_text("WEBVTT\n\n", encoding="utf-8")
    result = parse_transcript(vtt)
    assert result == ""
