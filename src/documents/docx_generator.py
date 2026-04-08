"""DOCX generation from structured content using python-docx."""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.core.audit import log_action

logger = logging.getLogger(__name__)

AMNESTY_YELLOW = RGBColor(0xFF, 0xFF, 0x00)
AMNESTY_BLACK = RGBColor(0x00, 0x00, 0x00)


def generate_docx(
    content: dict[str, Any],
    output_path: Path,
    template_path: Path | None = None,
    workflow: str = "docx_generator",
) -> Path:
    """Generate a DOCX document from structured content.

    Args:
        content: Dictionary with document structure:
            - title: str — Document title
            - subtitle: str — Optional subtitle
            - sections: list[dict] — Each with 'heading' and 'body' keys
            - metadata: dict — Optional metadata (author, date, etc.)
        output_path: Where to save the DOCX.
        template_path: Optional path to a .docx template to use as base.
        workflow: Workflow name for audit logging.

    Returns:
        Path to the generated DOCX.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if template_path and template_path.exists():
        doc = Document(str(template_path))
    else:
        doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title
    title_para = doc.add_heading(content["title"], level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if content.get("subtitle"):
        subtitle_para = doc.add_paragraph(content["subtitle"])
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacer

    # Sections
    for section in content.get("sections", []):
        if section.get("heading"):
            doc.add_heading(section["heading"], level=2)
        if section.get("body"):
            doc.add_paragraph(section["body"])

    # Metadata
    core_props = doc.core_properties
    if content.get("metadata"):
        meta = content["metadata"]
        if meta.get("author"):
            core_props.author = meta["author"]
        if meta.get("subject"):
            core_props.subject = meta["subject"]

    doc.save(str(output_path))

    log_action(
        workflow=workflow,
        action="docx_generated",
        actor="system",
        target=str(output_path),
        details={"title": content["title"]},
    )
    logger.info("Generated DOCX: %s", output_path)
    return output_path
